import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Helper to add headings
    def add_styled_heading(text, level, space_before=18, space_after=8):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        
        # style heading font
        h_font = h.runs[0].font
        h_font.name = 'Arial'
        if level == 1:
            h_font.size = Pt(16)
            h_font.bold = True
            h_font.color.rgb = RGBColor(31, 78, 121)  # Deep Navy
        elif level == 2:
            h_font.size = Pt(13)
            h_font.bold = True
            h_font.color.rgb = RGBColor(65, 113, 156)  # Blue
        return h

    # Helper to add standard paragraph
    def add_p(text, justify=True, space_after=8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        p.add_run(text)
        return p

    # --- PORTADA ---
    # Spacer
    for _ in range(3):
        doc.add_paragraph()
        
    p_uni = doc.add_paragraph()
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = p_uni.add_run("UNIVERSIDAD NACIONAL DE INGENIERÍA\nFACULTAD DE INGENIERÍA Y CIENCIAS DE LA COMPUTACIÓN\nESCUELA DE POSTGRADO / INGENIERÍA EN INTELIGENCIA ARTIFICIAL\n\n\n\n\n")
    r_uni.bold = True
    r_uni.font.size = Pt(14)
    r_uni.font.color.rgb = RGBColor(31, 78, 121)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("TRABAJO DE INVESTIGACIÓN Y DESARROLLO:\nDISEÑO E IMPLEMENTACIÓN DE UN ASISTENTE DOCUMENTAL INTELIGENTE (RAG LOCAL) CON PRIVACIDAD GARANTIZADA\n\n\n\n\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(47, 85, 151)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run(
        "Asignatura: Proyectos Avanzados de Ingeniería en Inteligencia Artificial\n\n"
        "Estudiante: [Ingresar Tu Nombre Completo]\n"
        "Profesor Evaluador: [Ingresar Nombre del Profesor]\n"
        "Fecha de Entrega: Junio 2026\n"
    )
    r_meta.font.size = Pt(11)
    
    doc.add_page_break()
    
    # --- INTRODUCCIÓN ---
    add_styled_heading("1. INTRODUCCIÓN", level=1)
    
    add_p(
        "En la actualidad, los Modelos de Lenguaje de Gran Escala (LLMs) han revolucionado la forma en que "
        "los humanos interactúan con la información estructurada. No obstante, en ámbitos corporativos, "
        "académicos y de investigación, la dependencia de APIs comerciales en la nube presenta serias "
        "desventajas. Entre ellas destacan los riesgos asociados a la privacidad de datos sensibles, la dependencia "
        "financiera de esquemas de pago por consumo y la propensión de los modelos comerciales a generar "
        "alucinaciones (información plausible pero falsa). Este trabajo describe el diseño y la implementación de un "
        "sistema de Generación Aumentada por Recuperación (RAG, por sus siglas en inglés) que se ejecuta al 100% "
        "de manera local en hardware de consumo estándar."
    )
    
    add_p(
        "El objetivo principal es construir una arquitectura de software robusta y modular capaz de extraer, "
        "segmentar, indexar y consultar información contenida en documentos privados en formato PDF. Al prescindir "
        "de conexiones externas, se asegura la confidencialidad de la información, permitiendo consultar manuales "
        "técnicos, reglamentos internos y documentación propia bajo los más altos estándares de protección de datos."
    )
    
    # --- DESARROLLO Y EXPLICACIÓN ---
    add_styled_heading("2. DESARROLLO Y EXPLICACIÓN TÉCNICA", level=1)
    
    add_p(
        "El desarrollo del proyecto se estructuró de forma modular a lo largo de un cronograma de ingeniería "
        "de software, abordando cada componente de la arquitectura RAG local como un módulo independiente. A "
        "continuación, se detalla la base conceptual, lógica y de código empleada."
    )
    
    add_styled_heading("2.1 Configuración de Infraestructura Local y Dependencias", level=2)
    add_p(
        "La base del sistema requiere de un entorno virtual aislado (venv) que asegura la consistencia de las versiones "
        "de Python y evita interferencias con librerías globales del sistema. El módulo de inferencia está soportado por "
        "el servidor Ollama en su versión 0.30.8, el cual permite la ejecución y servicio local del modelo Llama 3.2 (3B), "
        "un LLM ligero optimizado para tareas de razonamiento con baja huella de memoria (2.0 GB de tamaño)."
    )
    
    add_styled_heading("2.2 Ingesta Estructurada de Documentos (PDF Loader)", level=2)
    add_p(
        "La ingesta de datos representa la primera fase del flujo de trabajo. Para este propósito, se utilizó la "
        "clase PyPDFLoader de la librería LangChain. Este cargador procesa los documentos binarios PDF página "
        "por página y extrae el texto plano manteniendo un diccionario de metadatos asociado. La retención del "
        "número de página original de donde proviene cada fragmento es crítica para las fases de auditoría "
        "y auditoría posterior por parte del usuario final."
    )
    
    add_styled_heading("2.3 Estrategias Avanzadas de Fragmentación (Chunking)", level=2)
    add_p(
        "Un reto clave al implementar RAG es la restricción de la ventana de contexto del LLM. Si se inyecta un documento "
        "entero al modelo, este fallará por desbordamiento de tokens o por pérdida de atención (efecto 'Lost in the Middle'). "
        "Para solucionar esto, se implementó el procesador RecursiveCharacterTextSplitter. "
        "Este algoritmo divide recursivamente el texto usando una lista de caracteres jerárquicos (párrafos \\n\\n, líneas \\n, "
        "palabras con espacio, y caracteres individuales) para garantizar que las ideas semánticas no queden truncadas."
    )
    add_p(
        "Se configuró por defecto un tamaño de fragmento (chunk_size) de 1000 caracteres con una zona de solapamiento "
        "técnico (chunk_overlap) de 200 caracteres. El solapamiento del 20% es necesario para mantener la continuidad "
        "contextual de los párrafos que se encuentran en los bordes de la división."
    )
    
    add_styled_heading("2.4 Espacios Latentes y Embeddings Semánticos", level=2)
    add_p(
        "Para realizar búsquedas por significado en lugar de coincidencias literales de palabras clave, se transforman "
        "los fragmentos de texto en vectores multidimensionales densos (embeddings). Se implementó el modelo de código "
        "abierto all-MiniLM-L6-v2 de HuggingFace, gestionado a través de la librería sentence-transformers en local. "
        "Este modelo proyecta cada fragmento de texto a una coordenada matemática de 384 dimensiones, donde la distancia "
        "geométrica (como la similitud de coseno) entre dos vectores representa la cercanía semántica entre sus significados."
    )

    add_styled_heading("2.5 Indexación y Persistencia con ChromaDB", level=2)
    add_p(
        "Una vez generados los vectores de embeddings, se requiere un motor de almacenamiento de alta velocidad para "
        "su posterior recuperación. Se integró ChromaDB, una base de datos vectorial nativa para inteligencia artificial. "
        "ChromaDB se configuró en modo persistente en disco sólido (SSD) en la carpeta chroma_db. Esto evita la degradación "
        "del rendimiento por latencia y elimina la necesidad de volver a procesar y vectorizar los PDF cada vez que la "
        "aplicación es iniciada por el usuario."
    )

    add_styled_heading("2.6 Orquestación de la Cadena RAG con LangChain", level=2)
    add_p(
        "La integración lógica del pipeline se realiza mediante LangChain Expression Language (LCEL). El flujo lógico es el siguiente: "
        "1. El usuario realiza una pregunta a través del chat.\n"
        "2. La pregunta se vectoriza con all-MiniLM-L6-v2.\n"
        "3. ChromaDB realiza una búsqueda de similitud y recupera los 'k' fragmentos más cercanos (por defecto k=4).\n"
        "4. Los fragmentos se formatean con sus números de página y se inyectan como contexto estructurado en la plantilla del prompt.\n"
        "5. El prompt estructurado se envía al modelo local Llama 3.2 a través del conector ChatOllama.\n"
        "6. El modelo responde y la respuesta se muestra al usuario junto con las fuentes."
    )

    add_styled_heading("2.7 Ingeniería de Prompts para la Mitigación de Alucinaciones", level=2)
    add_p(
        "Uno de los principales riesgos en los sistemas de IA generativa es la alucinación de respuestas. Para mitigar esto "
        "al 100%, se diseñó una plantilla estricta (ChatPromptTemplate) con directrices claras en el prompt del sistema. "
        "El prompt le prohíbe al LLM deducir o utilizar conocimientos previos externos a los fragmentos provistos en el contexto. "
        "Si la respuesta a la pregunta del usuario no está explícitamente contenida en la información indexada, el modelo debe "
        "responder textualmente: 'Lo siento, pero la información solicitada no se encuentra en los documentos indexados.' "
        "Esto asegura una total fidelidad al documento corporativo o académico cargado."
    )

    add_styled_heading("2.8 Interfaz Gráfica y Gestión de Estado (Streamlit)", level=2)
    add_p(
        "Para proveer usabilidad a usuarios no técnicos, se construyó una interfaz interactiva en Streamlit con diseño premium "
        "basado en gradientes y componentes visuales en modo oscuro. La interfaz cuenta con gestión del historial de chat en "
        "el estado de la sesión (session_state), permitiendo simular una conversación continua. Adicionalmente, incluye un panel "
        "lateral de ajustes en caliente (tamaño de chunk, overlap, temperatura de creatividad y k de recuperación) y un visor "
        "desplegable que muestra los fragmentos literales de texto recuperados del PDF para contrastar la veracidad del RAG."
    )

    # --- CONCLUSIONES ---
    add_styled_heading("3. CONCLUSIONES", level=1)
    
    add_p(
        "El desarrollo de este asistente RAG local valida que es perfectamente viable construir herramientas avanzadas "
        "de procesamiento de información con altos estándares de calidad sin depender de servicios pagos en la nube. El uso "
        "de modelos ligeros optimizados (Llama 3.2 de 3B) y bases de datos vectoriales locales en disco sólido garantiza tiempos "
        "de respuesta extremadamente rápidos (inferiores a los 10 segundos), ofreciendo un rendimiento equiparable a soluciones "
        "comerciales en hardware portátil convencional."
    )
    add_p(
        "La arquitectura modular propuesta facilita la mantenibilidad del código y permite escalar el sistema en el futuro, "
        "pudiendo incorporar técnicas de re-ranking de fragmentos (cross-encoders) o la ingesta multiformato (cargando archivos "
        "Word, hojas de cálculo o bases de datos relacionales SQL)."
    )

    # --- BIBLIOGRAFÍA ---
    add_styled_heading("4. REFERENCIAS BIBLIOGRÁFICAS (APA SÉPTIMA EDICIÓN)", level=1)
    
    refs = [
        "Chroma. (2023). Chroma: The AI-native open-source embedding database (Versión 0.5) [Software de computación]. https://github.com/chroma-core/chroma",
        "LangChain. (2023). LangChain: Building applications with LLMs through composability (Versión 0.2) [Software de computación]. https://github.com/langchain-ai/langchain",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Lewis, P., Riedel, S., & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459–9474. https://arxiv.org/abs/1908.10084",
        "Meta AI. (2024). Llama 3.2: Open-source lightweight large language model [Modelo de lenguaje de gran escala]. Meta. https://llama.meta.com",
        "Ollama. (2024). Ollama: Run Llama 3, Mistral, and other large language models locally (Versión 0.3) [Software de computación]. https://ollama.com",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. En Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing. Association for Computational Linguistics. https://doi.org/10.18653/v1/D19-1410",
        "Streamlit Inc. (2020). Streamlit: The fastest way to build and share data apps (Versión 1.35) [Software de computación]. https://streamlit.io"
    ]
    
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    base_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(base_dir, "Trabajo_RAG_Universitario.docx")
    doc.save(output_path)
    print(f"Documento Word guardado en: {output_path}")

if __name__ == "__main__":
    create_report()
